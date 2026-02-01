import os
import sqlite3
from datetime import datetime
from io import BytesIO

from flask import (
    Flask, render_template, request, redirect, url_for, flash, send_file, jsonify
)

from docx import Document

# =====================================================
# CONFIGURAÇÕES
# =====================================================
APP_NAME = "EthosPsi"
app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-ethospsi-secret-final-v5")

DATA_DIR = os.path.abspath("./data")
os.makedirs(DATA_DIR, exist_ok=True)
DB_PATH = os.path.join(DATA_DIR, "ethospsi.sqlite3")

# =====================================================
# LINKS OFICIAIS
# =====================================================
LINKS_OFICIAIS = {
    "codigo_etica_pdf_2025": "https://site.cfp.org.br/wp-content/uploads/2012/07/codigo-de-etica-psicologia.pdf",
    "codigo_etica_pdf": "https://site.cfp.org.br/wp-content/uploads/2012/07/codigo-de-etica-psicologia.pdf",
    "tabela_honorarios_cfp": "https://site.cfp.org.br/servicos/tabela-de-honorarios/",
    "tabela_honorarios_pdf_ate_julho_2025": "https://site.cfp.org.br/wp-content/uploads/2025/12/3699.1___ANEXO_REF_AO_OFICIO_N__009_20225___FENAPSI.pdf",
}

# =====================================================
# TEXTO BASE (CÓDIGO DE ÉTICA RESUMIDO PARA BUSCA)
# =====================================================
TEXTO_CODIGO_ETICA = """
PRINCÍPIOS FUNDAMENTAIS
I. O psicólogo baseará o seu trabalho no respeito e na promoção da liberdade, da dignidade, da igualdade e da integridade do ser humano.
II. O psicólogo trabalhará visando promover a saúde e a qualidade de vida.

DAS RESPONSABILIDADES DO PSICÓLOGO - Art. 1º São deveres fundamentais:
c) Prestar serviços psicológicos de qualidade, utilizando princípios fundamentados na ciência psicológica, na ética e na legislação.
j) Ter, para com o trabalho dos psicólogos e de outros profissionais, respeito, consideração e solidariedade.

Art. 2º Ao psicólogo é vedado:
a) Praticar ou ser conivente com quaisquer atos que caracterizem negligência, discriminação, exploração, violência, crueldade ou opressão.
b) Induzir a convicções políticas, filosóficas, morais, ideológicas, religiosas, de orientação sexual ou a qualquer tipo de preconceito.
f) Prestar serviços ou vincular o título de psicólogo a serviços de atendimento psicológico cujos procedimentos, técnicas e meios não estejam regulamentados ou reconhecidos pela profissão.
j) Estabelecer com a pessoa atendida, familiar ou terceiro, relação que possa interferir negativamente nos objetivos do serviço prestado.
o) Receber, pagar remuneração ou porcentagem por encaminhamento de serviços.
q) Realizar diagnósticos, divulgar procedimentos ou apresentar resultados em meios de comunicação de forma a expor pessoas.

SIGILO PROFISSIONAL
Art. 9º - É dever do psicólogo respeitar o sigilo profissional a fim de proteger, por meio da confidencialidade, a intimidade das pessoas.
Art. 10 - Nas situações em que se configure conflito entre as exigências decorrentes do disposto no Art. 9º e as afirmações dos princípios fundamentais deste Código, excetuando-se os casos previstos em lei, o psicólogo poderá decidir pela quebra de sigilo, baseando sua decisão na busca do menor prejuízo.
Art. 13 - No atendimento à criança, ao adolescente ou ao interdito, deve ser comunicado aos responsáveis o estritamente essencial para se promoverem medidas em seu benefício.
"""

# =====================================================
# LISTA DE PERGUNTAS (ORDEM DE EXIBIÇÃO)
# =====================================================
QUICK_QUESTIONS = [
    "Até onde vai o sigilo?",
    "Quando posso quebrar o sigilo?",
    "Posso confirmar para alguém que a pessoa é minha paciente?",
    "Posso falar do caso com meu cônjuge ou amigo?",
    "Como agir se um familiar pede informações do paciente?",
    "Como agir se o paciente pede segredo absoluto?",
    "Até onde vai o sigilo em caso de crime?",
    "Sou obrigada a fazer anotações?",
    "O que é obrigatório eu anotar no prontuário?",
    "Paciente pediu para não registrar no prontuário",
    "O paciente pode pedir cópia do prontuário?",
    "Por quanto tempo devo guardar prontuários?",
    "Posso usar prontuários de forma digital?",
    "Posso usar IA para escrever prontuário?",
    "Posso emitir declaração de comparecimento?",
    "Posso emitir laudo psicológico para processo?",
    "Posso emitir relatório para escola?",
    "Posso colocar CID em relatório?",
    "Posso cobrar por relatório psicológico?",
    "Posso atender amigos?",
    "Posso atender familiares?",
    "Posso atender familiares de ex-pacientes?",
    "Posso ir a eventos sociais em que meu paciente esta?",
    "Posso seguir paciente no Instagram?",
    "Preciso de contrato para terapia online?",
    "Como garantir sigilo no atendimento online?",
    "Posso atender paciente dirigindo?",
    "Posso atender de graça?",
    "Posso divulgar o valor da sessão no Instagram?",
    "Como lidar com inadimplência?",
    "Existe cura gay?",
    "Posso orar com o paciente na sessão?",
    "Posso recusar atendimento por conflito de valores?",
    "Como encerrar terapia de forma ética?",
    "Quando devo encaminhar um paciente?",
    "Posso atender adolescente sem os pais saberem?",
    "Como agir em suspeita de violência (rede de proteção)?",
]

# =====================================================
# HELPERS DE RESPOSTA (HTML)
# =====================================================
def _html_escape(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def _make_answer(title: str, bullets: list[str], delicate: bool = True) -> str:
    """Gera o HTML da resposta com o alerta de 'Questão Delicada'."""
    warn = ""
    if delicate:
        warn = """
        <div class="alert-box warning">
          <strong>Questão delicada:</strong> se houver dúvida prática, consulte a COF do seu CRP e leve para supervisão.
        </div>
        """
    lis = "".join([f"<li>{_html_escape(b)}</li>" for b in bullets if (b or "").strip()])
    return f"""
    <div class="resposta-humanizada">
      <h3>{_html_escape(title)}</h3>
      {warn}
      <ul>{lis}</ul>
    </div>
    """

# =====================================================
# BANCO DE DADOS DE RESPOSTAS (DICT)
# =====================================================
RESPOSTAS_DB = {
    # --- SIGILO ---
    "Até onde vai o sigilo?": _make_answer(
        "O sigilo é inerente à profissão (Art. 9º do CEpp).",
        [
            "Protege a intimidade da pessoa, grupos ou organizações.",
            "Não é absoluto: pode ser quebrado em situações de risco grave à vida (suicídio/homicídio), violência contra vulneráveis ou ordem judicial específica.",
            "Mesmo na quebra, deve-se restringir as informações ao estritamente necessário."
        ]
    ),
    "Quando posso quebrar o sigilo?": _make_answer(
        "Situações de exceção (Art. 10º do CEpp).",
        [
            "Quando houver risco grave e iminente à vida (da pessoa ou terceiros).",
            "Casos de violência contra criança, adolescente ou idoso (comunicação aos órgãos competentes).",
            "Decisão judicial fundamentada (embora o psicólogo possa arguir sigilo se julgar prejudicial).",
            "Em todos os casos, prestar apenas as informações estritamente necessárias."
        ]
    ),
    "Posso confirmar para alguém que a pessoa é minha paciente?": _make_answer(
        "Não. A própria existência do vínculo é sigilosa.",
        [
            "Confirmar o atendimento expõe a pessoa e viola o Art. 9º.",
            "Resposta padrão: 'Por questões éticas, não posso confirmar nem negar atendimento a qualquer pessoa'.",
            "Exceção: Se o paciente autorizar formalmente ou se for responsável legal de menor."
        ]
    ),
    "Posso falar do caso com meu cônjuge ou amigo?": _make_answer(
        "Não. Violação gravíssima de sigilo.",
        [
            "O Art. 9º veda a exposição da intimidade.",
            "Mesmo trocando nomes, detalhes podem identificar o paciente.",
            "Angústias do terapeuta devem ser tratadas em Supervisão ou Terapia Pessoal, nunca em conversas sociais."
        ]
    ),
    "Como agir se um familiar pede informações do paciente?": _make_answer(
        "Proteja o sigilo e o vínculo.",
        [
            "Para pacientes adultos: Não passe informações sem consentimento expresso.",
            "Para crianças/adolescentes: Pais têm direito a feedback, mas não ao conteúdo detalhado das sessões (Art. 13º). Informe apenas o necessário para promover medidas em benefício.",
            "Acolha a angústia da família, mas reafirme a ética."
        ]
    ),
    "Como agir se o paciente pede segredo absoluto?": _make_answer(
        "Alinhe expectativas (Contrato Terapêutico).",
        [
            "Explique que o sigilo é a regra, mas a lei obriga a quebra em risco de vida ou violência.",
            "Isso constrói confiança e transparência desde o início.",
            "Garanta que nada será revelado 'sem querer' ou por descuido."
        ]
    ),
    "Até onde vai o sigilo em caso de crime?": _make_answer(
        "O psicólogo não é investigador de polícia.",
        [
            "Crimes passados relatados em sessão (ex: roubo) estão sob sigilo.",
            "Crimes em andamento ou futuros com risco à vida (ex: planejamento de homicídio/suicídio) ou contra vulneráveis exigem quebra de sigilo para proteção (Art. 10º).",
            "Em dúvida, consulte o jurídico do CRP."
        ]
    ),

    # --- PRONTUÁRIO / DOCUMENTOS ---
    "Sou obrigada a fazer anotações?": _make_answer(
        "Sim. O registro documental é obrigatório.",
        [
            "Resolução CFP 01/2009: É dever do psicólogo manter registro documental.",
            "Serve para garantia de direitos do usuário, defesa do profissional e evolução do caso.",
            "A falta de registro é infração ética frequente em fiscalizações."
        ]
    ),
    "O que é obrigatório eu anotar no prontuário?": _make_answer(
        "Conteúdo mínimo (Res. CFP 01/2009).",
        [
            "Identificação do usuário.",
            "Avaliação da demanda e definição de objetivos.",
            "Registro da evolução (datas, procedimentos, síntese do atendimento).",
            "Encaminhamentos ou encerramento.",
            "Não é necessário transcrever a sessão inteira (diário), apenas a síntese técnica."
        ]
    ),
    "Paciente pediu para não registrar no prontuário": _make_answer(
        "O registro é dever do psicólogo, não escolha do paciente.",
        [
            "Explique que é uma obrigação legal (Res. 01/2009).",
            "Negocie o teor: você pode registrar de forma mais sintética, protegendo detalhes muito íntimos, mas mantendo a evolução técnica.",
            "O prontuário pertence ao paciente, mas a guarda é do psicólogo."
        ]
    ),
    "O paciente pode pedir cópia do prontuário?": _make_answer(
        "Sim. É direito do usuário (Res. CFP 01/2009).",
        [
            "O paciente tem acesso integral às suas informações.",
            "O psicólogo deve fornecer cópia ou acesso quando solicitado.",
            "Se houver risco de o conteúdo causar dano (ex: surto psicótico ao ler), o profissional deve oferecer acompanhamento ou entrevista devolutiva para explicar o conteúdo."
        ]
    ),
    "Por quanto tempo devo guardar prontuários?": _make_answer(
        "Mínimo de 05 anos.",
        [
            "Resolução CFP 01/2009: Guarda mínima de 5 anos após o último atendimento.",
            "Após esse prazo, podem ser incinerados ou destruídos de forma segura.",
            "Para crianças, recomenda-se guardar até a maioridade (precaução jurídica)."
        ]
    ),
    "Posso usar prontuários de forma digital?": _make_answer(
        "Sim, com requisitos de segurança.",
        [
            "O sistema deve garantir sigilo, autenticidade e integridade.",
            "Recomendado uso de Certificado Digital (ICP-Brasil) para assinatura.",
            "Evite guardar em Word/Excel sem senha ou em nuvens públicas não seguras."
        ]
    ),
    "Posso usar IA para escrever prontuário?": _make_answer(
        "Extremo cuidado. Risco de violação de sigilo.",
        [
            "Não insira nomes ou dados identificáveis em IAs públicas (ChatGPT, Gemini, etc.), pois os dados podem ser tratados fora do seu controle.",
            "A responsabilidade técnica do texto é 100% do psicólogo.",
            "O uso deve ser apenas para auxílio na redação, nunca para análise clínica automática."
        ]
    ),

    # --- DOCUMENTOS (Resolução 06/2019) ---
    "Posso emitir laudo psicológico para processo?": _make_answer(
        "Sim, se houver demanda e capacitação.",
        [
            "Deve seguir rigorosamente a Resolução CFP 06/2019.",
            "O Laudo é resultado de Avaliação Psicológica. Não emita laudo apenas com base em psicoterapia.",
            "Deve ser imparcial, objetivo e responder aos quesitos ou demanda específica."
        ]
    ),
    "Posso emitir declaração de comparecimento?": _make_answer(
        "Sim. É o documento mais simples.",
        [
            "Resolução 06/2019: Atesta apenas o comparecimento (dia, hora, duração).",
            "Não deve conter diagnóstico (CID) ou sintomas, a menos que estritamente necessário e solicitado.",
            "Serve para justificar falta no trabalho/escola."
        ]
    ),
    "Posso emitir relatório para escola?": _make_answer(
        "Sim, focado no processo de aprendizagem.",
        [
            "Não exponha a intimidade familiar para a escola.",
            "O foco deve ser: como as questões emocionais impactam a aprendizagem ou comportamento escolar.",
            "Sempre peça autorização dos responsáveis e, se possível, mostre o documento a eles antes de enviar."
        ]
    ),
    "Posso colocar CID em relatório?": _make_answer(
        "Apenas com autorização e se tecnicamente justificado.",
        [
            "O diagnóstico pertence ao paciente. Só coloque CID se o paciente solicitar ou autorizar.",
            "Em documentos para planos de saúde ou INSS, geralmente é exigido, mas discuta com o paciente antes.",
            "Evite rotulação desnecessária."
        ]
    ),
    "Posso cobrar por relatório psicológico?": _make_answer(
        "Depende do contexto.",
        [
            "Se for um relatório simples de evolução do tratamento, geralmente entende-se incluso no serviço.",
            "Se for um Laudo ou Avaliação Psicológica extra (documento complexo), pode ser cobrado à parte, desde que acordado previamente no Contrato."
        ]
    ),

    # --- RELACIONAMENTOS DUAIS ---
    "Posso atender amigos?": _make_answer(
        "Não. Veda relação que interfira na objetividade.",
        [
            "Art. 2º, j: Vedado estabelecer relação que possa interferir negativamente.",
            "A intimidade prévia contamina a transferência e a neutralidade técnica.",
            "Encaminhe para um colega de confiança."
        ]
    ),
    "Posso atender familiares?": _make_answer(
        "Não. Configura relação dual/múltipla.",
        [
            "Mesma lógica dos amigos: falta de isenção e risco de confusão de papéis.",
            "O vínculo pessoal pré-existente impede o vínculo profissional ético.",
            "Encaminhamento é a conduta correta."
        ]
    ),
    "Posso atender familiares de ex-pacientes?": _make_answer(
        "Cuidado. Avalie caso a caso.",
        [
            "Não é explicitamente proibido, mas pode gerar conflito se as histórias se cruzarem.",
            "Se o ex-paciente foi atendido há pouco tempo ou se o vínculo familiar é muito próximo, melhor evitar.",
            "Priorize a qualidade do serviço e o sigilo."
        ]
    ),
    "Posso ir a eventos sociais em que meu paciente esta?": _make_answer(
        "Situação delicada. Preserve o enquadre.",
        [
            "Se for evento grande (show, palestra), ok. Se for íntimo (aniversário na casa de amigo comum), evite.",
            "Se encontrar: cumprimente discretamente, não puxe assunto terapêutico.",
            "Proteja o sigilo: não deixe transparecer para outros que é seu paciente."
        ]
    ),
    "Posso seguir paciente no Instagram?": _make_answer(
        "Não recomendado. Proteja o setting.",
        [
            "Ter acesso à vida pessoal do paciente fora da sessão pode enviesar a escuta.",
            "O paciente ter acesso à sua vida pessoal pode interferir na transferência.",
            "Perfis profissionais são ok, mas evite seguir de volta (seguir o paciente) para manter a assimetria da relação."
        ]
    ),

    # --- ONLINE / TECNOLOGIA ---
    "Preciso de contrato para terapia online?": _make_answer(
        "Altamente recomendado.",
        [
            "Estabeleça regras claras sobre plataforma, falhas de conexão, privacidade do ambiente e pagamentos.",
            "Define o que acontece se a internet cair (remarca? cobra?).",
            "Protege ambas as partes."
        ]
    ),
    "Como garantir sigilo no atendimento online?": _make_answer(
        "Medidas técnicas e ambientais.",
        [
            "Use fones de ouvido.",
            "Esteja em sala fechada e isolada acusticamente.",
            "Peça para o paciente fazer o mesmo (garantir que ele esteja sozinho).",
            "Evite gravar sessões sem necessidade extrema e consentimento."
        ]
    ),
    "Posso atender paciente dirigindo?": _make_answer(
        "Não. Risco à segurança e falta de foco.",
        [
            "A sessão exige atenção plena e ambiente seguro.",
            "Dirigir exige atenção ao trânsito. Fazer os dois coloca o paciente em risco físico.",
            "Interrompa a sessão e peça para ele estacionar ou remarcar."
        ]
    ),

    # --- HONORÁRIOS ---
    "Posso atender de graça?": _make_answer(
        "Sim, mas com ética (não promocional).",
        [
            "O atendimento voluntário é permitido e nobre.",
            "Não pode ser usado para captar clientela (ex: '1ª sessão grátis' como marketing).",
            "Deve ser um trabalho social genuíno ou vinculado a instituição."
        ]
    ),
    "Posso divulgar o valor da sessão no Instagram?": _make_answer(
        "Não. Evite propaganda por preço.",
        [
            "O preço não deve ser o diferencial competitivo.",
            "Informe o valor apenas quando o interessado entrar em contato (direct/whatsapp).",
            "Se houver dúvidas específicas, consulte orientações do CRP/CFP e pratique comunicação responsável."
        ]
    ),
    "Como lidar com inadimplência?": _make_answer(
        "Sem exposição e com diálogo.",
        [
            "O psicólogo não pode expor o paciente a situações vexatórias de cobrança.",
            "Tente renegociar, parcelar ou entender o motivo.",
            "Se a inadimplência persistir, pode-se suspender o atendimento, com encaminhamento e encerramento ético."
        ]
    ),

    # --- TEMAS SENSÍVEIS / RELIGIÃO / SEXUALIDADE ---
    "Existe cura gay?": _make_answer(
        "Não. E é proibido oferecer.",
        [
            "A homossexualidade não é doença, perversão ou distúrbio.",
            "É vedado colaborar com serviços que proponham 'tratamento' ou 'cura' da homossexualidade.",
            "O foco ético é acolher sofrimento, principalmente o decorrente do preconceito."
        ]
    ),
    "Posso orar com o paciente na sessão?": _make_answer(
        "Evite misturar técnica com prática religiosa.",
        [
            "A psicologia é laica.",
            "O psicólogo deve respeitar a crença do paciente, mas não deve induzir práticas religiosas durante o atendimento técnico.",
            "Se espiritualidade for tema do paciente, pode ser acolhida como conteúdo, sem ritualização."
        ]
    ),
    "Posso recusar atendimento por conflito de valores?": _make_answer(
        "Sim. É ético reconhecer limites.",
        [
            "Se uma demanda impedir uma escuta responsável, recusar pode ser o mais ético.",
            "Justifique de forma respeitosa e encaminhe para outro profissional qualificado.",
            "Evite moralização ou tentativa de 'corrigir' o paciente."
        ]
    ),

    # --- ENCERRAMENTO E ENCAMINHAMENTO ---
    "Como encerrar terapia de forma ética?": _make_answer(
        "Planejamento e Autonomia.",
        [
            "O encerramento (alta) deve ser trabalhado processualmente, não abruptamente.",
            "Deve visar a autonomia do paciente.",
            "Se for interrupção (pelo terapeuta), oferecer encaminhamento e apoiar continuidade do cuidado."
        ]
    ),
    "Quando devo encaminhar um paciente?": _make_answer(
        "Limites técnicos ou pessoais.",
        [
            "Quando a demanda exige competência técnica que você não possui.",
            "Quando há quebra do vínculo de confiança ou conflito de valores intransponível.",
            "Faça o encaminhamento de forma responsável, indicando serviços adequados."
        ]
    ),
    "Posso atender adolescente sem os pais saberem?": _make_answer(
        "Depende da situação e do contrato.",
        [
            "Adolescente tem direito a escuta e sigilo em muitos contextos.",
            "Para continuidade e responsabilidade legal/financeira, responsáveis normalmente precisam estar cientes.",
            "Em suspeita de violência intrafamiliar, priorize proteção e acione rede (Conselho Tutelar), seguindo o mínimo necessário."
        ]
    ),
    "Como agir em suspeita de violência (rede de proteção)?": _make_answer(
        "Notificação e Proteção.",
        [
            "Em casos de violência contra criança, adolescente, idoso ou mulher, pode haver dever de notificação conforme legislação aplicável.",
            "Não confronte o suposto agressor se isso colocar a vítima em risco.",
            "Acione a rede de proteção (CREAS, Conselho Tutelar, Delegacia) de forma articulada."
        ]
    ),
}

# =====================================================
# GERAÇÃO DE RESPOSTAS
# =====================================================
def generate_answer_for_question(q: str) -> str:
    """Retorna a resposta específica do DB ou um fallback genérico."""
    if q in RESPOSTAS_DB:
        return RESPOSTAS_DB[q]

    return _make_answer(
        "Consulte o Código de Ética",
        [
            "Esta pergunta requer análise específica dos artigos e resoluções aplicáveis.",
            "Recomenda-se leitura da normativa pertinente ao tema (ex: documentos, online, registro).",
            "Na dúvida, consulte a COF do seu CRP e leve para supervisão."
        ],
        delicate=True
    )

# =====================================================
# BANCO DE DADOS (SQLITE)
# =====================================================
def db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = db()
    cur = conn.cursor()
    cur.execute("""CREATE TABLE IF NOT EXISTS documents (id INTEGER PRIMARY KEY, title TEXT, created_at TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS chunks (id INTEGER PRIMARY KEY, doc_id INTEGER, chunk_text TEXT)""")
    cur.execute("""CREATE TABLE IF NOT EXISTS qa_history (id INTEGER PRIMARY KEY, question TEXT, answer TEXT, created_at TEXT)""")
    conn.commit()
    conn.close()

def clear_documents():
    conn = db()
    conn.execute("DELETE FROM chunks")
    conn.execute("DELETE FROM documents")
    conn.commit()
    conn.close()

def save_history(question: str, answer: str):
    conn = db()
    conn.execute(
        "INSERT INTO qa_history (question, answer, created_at) VALUES (?,?,?)",
        (question, answer, datetime.now().strftime("%d/%m %H:%M"))
    )
    conn.commit()
    conn.close()

def get_history(limit: int = 50):
    conn = db()
    rows = conn.execute("SELECT * FROM qa_history ORDER BY id DESC LIMIT ?", (limit,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def stats():
    conn = db()
    try:
        d = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
        c = conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
        h = conn.execute("SELECT COUNT(*) FROM qa_history").fetchone()[0]
    except Exception:
        conn.close()
        return {"documents": 0, "chunks": 0, "history": 0}
    conn.close()
    return {"documents": d, "chunks": c, "history": h}

# =====================================================
# INDEX e BUSCA (MANTIDOS PARA POSSÍVEL USO FUTURO)
# =====================================================
def index_content(title: str, text: str):
    chunks = [c.strip() for c in text.split('\n') if len(c.strip()) > 20]
    conn = db()
    cur = conn.cursor()
    cur.execute("INSERT INTO documents (title, created_at) VALUES (?,?)", (title, datetime.now().strftime("%Y-%m-%d")))
    doc_id = cur.lastrowid
    for c in chunks:
        cur.execute("INSERT INTO chunks (doc_id, chunk_text) VALUES (?,?)", (doc_id, c))
    conn.commit()
    conn.close()

def simple_search(query: str):
    conn = db()
    terms = (query or "").lower().split()
    keywords = [t for t in terms if len(t) > 3]
    if not keywords:
        return []

    sql = "SELECT chunk_text FROM chunks WHERE " + " OR ".join(["chunk_text LIKE ?"] * len(keywords))
    params = [f"%{k}%" for k in keywords]

    rows = conn.execute(sql, params).fetchall()
    conn.close()

    seen = set()
    unique_rows = []
    for r in rows:
        if r[0] not in seen:
            unique_rows.append(r[0])
            seen.add(r[0])
    return unique_rows[:3]

# =====================================================
# FERRAMENTAS (CONTRATO, HONORARIOS, POLITICAS, REDE)
# =====================================================
def gerar_contrato_texto(data: dict) -> str:
    modalidade = data.get("modalidade", "Online")
    plataforma = data.get("plataforma", "Google Meet")
    duracao = data.get("duracao", "50")
    frequencia = data.get("frequencia", "semanal")
    canal = data.get("canal", "WhatsApp")
    prazo_cancel = data.get("prazo_cancel", "24")
    falta_cobra = data.get("falta_cobra", "sim")
    atraso = data.get("atraso", "15")
    queda = data.get("queda", "10")
    pagamento = data.get("pagamento", "pix")
    recibo = data.get("recibo", "sim")
    reembolso = data.get("reembolso", "nao")
    emergencias = data.get("emergencias", "sim")
    sigilo = data.get("sigilo", "sim")
    grava = data.get("grava", "nao")

    detalhe_modalidade = (
        "Atendimento presencial em ambiente privativo, com início e término conforme horário agendado."
        if modalidade.lower() == "presencial"
        else f"Atendimento online por {plataforma}, com orientações de privacidade (local reservado e, se possível, uso de fone)."
    )
    falta_txt = "Sessões não desmarcadas dentro do prazo são cobradas." if falta_cobra == "sim" else "Sessões não desmarcadas dentro do prazo podem ser remanejadas conforme disponibilidade e critério."
    recibo_txt = "Recibo pode ser emitido mediante solicitação." if recibo == "sim" else "Recibo não é emitido."
    reembolso_txt = "Em caso de interrupção do serviço, valores antecipados podem ser ajustados conforme sessões realizadas." if reembolso == "sim" else "Não há reembolso automático para faltas ou cancelamentos fora do prazo."
    emerg_txt = "Este serviço não é plantão de urgência. Em risco imediato, recomenda-se acionar rede de apoio e serviços locais." if emergencias == "sim" else "Este serviço não realiza atendimentos de urgência."
    sig_txt = "O sigilo profissional é regra. Exceções são raras e seguem princípio do mínimo necessário." if sigilo == "sim" else "O sigilo profissional orienta a prática, com cuidado especial para privacidade."
    grava_txt = "Gravações não são permitidas sem consentimento explícito das partes e finalidade justificada." if grava == "nao" else "Gravações podem ocorrer apenas com consentimento explícito e acordo sobre guarda e acesso."

    return f"""CONTRATO TERAPÊUTICO (MODELO)

1) Modalidade e setting
- Modalidade: {modalidade}
- {detalhe_modalidade}

2) Duração e frequência
- Duração média da sessão: {duracao} minutos
- Frequência sugerida: {frequencia}

3) Comunicação fora da sessão
- Canal de contato: {canal}
- Finalidade: logística (remarcação, confirmação e avisos)
- Mensagens longas são preferencialmente tratadas em sessão.

4) Cancelamentos, faltas e atrasos
- Prazo para desmarcação: {prazo_cancel} horas
- Tolerância de atraso: {atraso} minutos (respeitando o horário final)
- Faltas: {falta_txt}

5) Atendimento online e queda de conexão (se aplicável)
- Em queda: aguardar {queda} minutos e tentar reconectar
- Se não retomar: registrar tentativa e remarcar conforme política.

6) Pagamento e recibos
- Forma de pagamento: {pagamento}
- {recibo_txt}
- {reembolso_txt}

7) Sigilo e privacidade
- {sig_txt}

8) Gravações
- {grava_txt}

9) Limites e emergências
- {emerg_txt}

10) Encerramento
- Encerramento por alta, acordo, limites de agenda ou indicação clínica.
- Quando possível, será trabalhado em sessão, com orientações e encaminhamentos.

Observação
Este documento é um modelo informacional e pode ser adaptado conforme contexto e critérios profissionais.
"""

def calc_honorarios(d: dict) -> dict:
    custos_fixos = float(d.get("custos_fixos", 0) or 0)
    custos_variaveis_mes = float(d.get("custos_variaveis_mes", 0) or 0)
    pro_labore = float(d.get("pro_labore", 0) or 0)
    impostos_perc = float(d.get("impostos_perc", 0) or 0) / 100.0
    semanas_mes = float(d.get("semanas_mes", 4.3) or 4.3)

    sessoes_semana = float(d.get("sessoes_semana", 0) or 0)
    duracao_min = float(d.get("duracao_min", 50) or 50)
    admin_min = float(d.get("admin_min", 10) or 10)
    faltas_perc = float(d.get("faltas_perc", 0) or 0) / 100.0

    custo_total_mes = custos_fixos + custos_variaveis_mes + pro_labore
    sessoes_mes_brutas = sessoes_semana * semanas_mes
    sessoes_mes_liquidas = max(0.0, sessoes_mes_brutas * (1.0 - faltas_perc))

    if sessoes_mes_liquidas <= 0:
        return {"ok": False, "erro": "Defina sessões por semana e faltas em um valor que gere ao menos 1 sessão/mês efetiva."}

    if impostos_perc >= 0.95:
        impostos_perc = 0.95

    receita_bruta_necessaria = custo_total_mes / max(0.01, (1.0 - impostos_perc))
    preco_min_sessao = receita_bruta_necessaria / sessoes_mes_liquidas

    tempo_por_sessao_min = duracao_min + admin_min
    horas_trabalho_mes = (tempo_por_sessao_min * sessoes_mes_brutas) / 60.0
    if horas_trabalho_mes <= 0:
        horas_trabalho_mes = 0.1

    receita_por_hora_bruta = (preco_min_sessao * sessoes_mes_liquidas) / horas_trabalho_mes

    return {
        "ok": True,
        "custo_total_mes": round(custo_total_mes, 2),
        "receita_bruta_necessaria": round(receita_bruta_necessaria, 2),
        "sessoes_mes_brutas": round(sessoes_mes_brutas, 1),
        "sessoes_mes_liquidas": round(sessoes_mes_liquidas, 1),
        "preco_min_sessao": round(preco_min_sessao, 2),
        "tempo_por_sessao_min": round(tempo_por_sessao_min, 0),
        "horas_trabalho_mes": round(horas_trabalho_mes, 1),
        "receita_por_hora_bruta": round(receita_por_hora_bruta, 2),
    }

def gerar_politica(data: dict) -> dict:
    tipo = data.get("tipo", "faltas")
    modalidade = data.get("modalidade", "Online")
    prazo = data.get("prazo", "24")
    atraso = data.get("atraso", "15")
    canal = data.get("canal", "WhatsApp")
    falta_cobra = data.get("falta_cobra", "sim")
    reembolso = data.get("reembolso", "nao")
    mensagens = data.get("mensagens", "logistica")
    queda = data.get("queda", "10")
    pagamento = data.get("pagamento", "pix")

    base_header = "POLÍTICA (TEXTO PRONTO PARA COPIAR)\n\n"

    if tipo == "faltas":
        texto = f"""{base_header}Política de cancelamento e faltas

- Prazo para desmarcação: {prazo} horas.
- Atrasos: tolerância de {atraso} minutos, respeitando o horário final.
- Falta sem aviso ou cancelamento fora do prazo: {"sessão é cobrada" if falta_cobra == "sim" else "pode ser remanejada conforme disponibilidade e critério"}.
- Canal para desmarcação: {canal}.
"""
        return {"titulo": "Faltas e cancelamentos", "texto": texto}

    if tipo == "mensagens":
        if mensagens == "logistica":
            regra = "Mensagens são destinadas apenas à logística (remarcação, confirmação e avisos)."
        elif mensagens == "curtas":
            regra = "Mensagens devem ser curtas e objetivas. Conteúdos terapêuticos serão priorizados em sessão."
        else:
            regra = "Mensagens não substituem a sessão. Em caso de necessidade importante, combinaremos o melhor encaminhamento."

        texto = f"""{base_header}Política de mensagens e contato fora da sessão

- Canal principal: {canal}.
- {regra}
- Este serviço não funciona como plantão de urgência.
"""
        return {"titulo": "Mensagens e contato", "texto": texto}

    if tipo == "reembolso":
        texto = f"""{base_header}Política de pagamentos e reembolso

- Forma de pagamento: {pagamento}.
- Reembolso: {"pode haver ajuste proporcional em caso de interrupção do serviço, conforme sessões realizadas" if reembolso == "sim" else "não há reembolso automático para faltas ou cancelamentos fora do prazo"}.
"""
        return {"titulo": "Pagamentos e reembolso", "texto": texto}

    if tipo == "online":
        texto = f"""{base_header}Protocolo de atendimento online

- Modalidade: {modalidade}.
- Recomenda-se ambiente privado e uso de fone.
- Em queda de conexão: aguardar {queda} minutos e tentar reconectar.
- Se não retomar: confirmar por {canal} e remarcar conforme disponibilidade.
"""
        return {"titulo": "Atendimento online", "texto": texto}

    if tipo == "sigilo":
        texto = f"""{base_header}Política de sigilo e privacidade

- O sigilo profissional é regra e protege a intimidade e o vínculo terapêutico.
- Informações só são compartilhadas em situações excepcionais, seguindo o princípio do mínimo necessário.
"""
        return {"titulo": "Sigilo e privacidade", "texto": texto}

    return {"titulo": "Política", "texto": f"{base_header}Escolha uma política para gerar um texto pronto."}

def gerar_rede(data: dict) -> dict:
    destino = data.get("destino", "psiquiatria")
    canal = data.get("canal", "WhatsApp")
    inclui_autorizacao = data.get("autorizacao", "sim")

    autorizacao_txt = (
        "Antes de qualquer contato com terceiros, solicite autorização por escrito do paciente (ou responsável legal), delimitando o que pode ser compartilhado e com qual finalidade.\n\n"
        if inclui_autorizacao == "sim" else ""
    )

    if destino == "psiquiatria":
        texto = f"""ROTEIRO DE REDE: Psiquiatria

{autorizacao_txt}Mensagem para encaminhamento
Canal sugerido: {canal}

Olá, tudo bem?
Sou psicóloga e estou acompanhando a pessoa em psicoterapia. Com autorização expressa, gostaria de encaminhar para avaliação psiquiátrica.
"""
        return {"titulo": "Encaminhamento para Psiquiatria", "texto": texto}

    if destino == "autorizacao":
        texto = """MODELO: Autorização para contato com terceiros

Eu, ______________________________, autorizo a psicóloga ______________________________ (CRP ________) a realizar contato profissional com ______________________________, com a finalidade de ______________________________.
"""
        return {"titulo": "Autorização por escrito", "texto": texto}

    return {"titulo": "Rede", "texto": "Escolha um destino para gerar um roteiro."}

# =====================================================
# DOCX DOWNLOAD
# =====================================================
def _sanitize_filename(name: str) -> str:
    keep = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789-_ "
    cleaned = "".join([c if c in keep else "_" for c in (name or "")]).strip()
    return cleaned[:80] if cleaned else "documento"

def _make_docx_bytes(title: str, text: str) -> BytesIO:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    lines = (text or "").replace("\r\n", "\n").split("\n")
    for line in lines:
        doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

@app.route("/download-docx", methods=["POST"])
def download_docx():
    title = (request.form.get("doc_title") or "Documento").strip()
    text = request.form.get("doc_text") or ""
    filename = _sanitize_filename(request.form.get("doc_filename") or title)

    if not text.strip():
        flash("Nada para baixar. Gere o documento primeiro.", "success")
        return redirect(request.referrer or url_for("home"))

    bio = _make_docx_bytes(title=title, text=text)
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"{filename}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# =====================================================
# ROTAS
# =====================================================
@app.route("/", methods=["GET", "POST"])
def home():
    answer = None

    if request.method == "POST":
        if "load_bases" in request.form:
            clear_documents()
            index_content("Código de Ética (Resumo)", TEXTO_CODIGO_ETICA)
            flash("Base atualizada com sucesso!", "success")
            return redirect(url_for("home"))

        q = (request.form.get("q") or "").strip()
        if q:
            answer = generate_answer_for_question(q)
            save_history(q, answer)

    all_questions = [{"text": q} for q in QUICK_QUESTIONS]

    return render_template(
        "home.html",
        app_name=APP_NAME,
        stats=stats(),
        history=get_history(50),
        answer=answer,
        questions=all_questions,
    )

# Rota opcional: para o front abrir modal via fetch (sem rolar)
@app.route("/qa", methods=["GET"])
def qa_get():
    q = (request.args.get("q") or "").strip()
    if not q:
        return jsonify({"ok": False, "error": "missing q"}), 400
    html = generate_answer_for_question(q)
    # não salva no histórico aqui, porque só “abrir” não significa que perguntou;
    # o template pode optar por chamar /qa e também postar o form se quiser.
    return jsonify({"ok": True, "question": q, "answer_html": html})

@app.route("/recursos")
def recursos():
    return render_template("resources.html", app_name=APP_NAME, links=LINKS_OFICIAIS)

@app.route("/contrato", methods=["GET", "POST"])
def contrato():
    contrato_txt = None
    if request.method == "POST":
        contrato_txt = gerar_contrato_texto(request.form)
    return render_template("contrato.html", app_name=APP_NAME, contrato_txt=contrato_txt)

@app.route("/honorarios", methods=["GET", "POST"])
def honorarios():
    resultado = None
    if request.method == "POST":
        resultado = calc_honorarios(request.form)
    return render_template("honorarios.html", app_name=APP_NAME, resultado=resultado, links=LINKS_OFICIAIS)

@app.route("/politicas", methods=["GET", "POST"])
def politicas():
    out = None
    if request.method == "POST":
        out = gerar_politica(request.form)
    return render_template("politicas.html", app_name=APP_NAME, out=out)

@app.route("/rede", methods=["GET", "POST"])
def rede():
    out = None
    if request.method == "POST":
        out = gerar_rede(request.form)
    return render_template("rede.html", app_name=APP_NAME, out=out)

@app.route("/admin")
def admin():
    return render_template("admin.html", stats=stats(), app_name=APP_NAME)

# =====================================================
# INICIALIZAÇÃO
# =====================================================
if __name__ == "__main__":
    init_db()
    if stats()["chunks"] == 0:
        index_content("Código de Ética (Resumo)", TEXTO_CODIGO_ETICA)
    app.run(debug=True, port=5000)
